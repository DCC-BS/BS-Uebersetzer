from typing import Optional, Union
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
from tqdm import tqdm
import logging
import contextlib
from contextvars import ContextVar
from utils import detect_language, split_into_sentences
from docx import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

tqdm_disabled = ContextVar("tqdm_disabled", default=False)


@contextlib.contextmanager
def suppress_output():
    """Context manager to temporarily suppress tqdm progress bars and logging."""
    original_log_level = logging.getLogger().level
    t = tqdm_disabled.set(True)
    try:
        logging.getLogger().setLevel(logging.CRITICAL + 1)
        yield
    finally:
        logging.getLogger().setLevel(original_log_level)
        tqdm_disabled.reset(t)


def translate_sentence(sentence, model_name):
    """Translate a single sentence using the specified Opus MT model."""
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    translator = pipeline(
        "translation", model=model, tokenizer=tokenizer, device="cuda"
    )
    translated_text = translator(sentence)[0]["translation_text"]
    return translated_text


def translate_sentences(sentences, target_language):
    """Translate all sentences with appropriate Opus MT model."""
    translated_sentences = []
    logger.info(f"Translating {len(sentences)} sentences to {target_language}")
    model_name = None

    with tqdm(
        total=len(sentences), desc="Translating sentences", disable=tqdm_disabled.get()
    ) as pbar:
        for sentence in sentences:
            source_language = detect_language(sentence)

            if source_language == "en":
                model_name = "Helsinki-NLP/opus-mt-en-de"
            elif source_language == "fr":
                model_name = "Helsinki-NLP/opus-mt-fr-de"
            elif source_language == "it":
                model_name = "Helsinki-NLP/opus-mt-it-de"
            else:
                logger.error(
                    f"Unsupported language detected: {source_language}, Text: {sentence}. Skipping this sentence."
                )
                translated_sentences.append(sentence)
                pbar.update(1)
                continue

            translated_sentence = translate_sentence(sentence, model_name)
            if translated_sentence:
                translated_sentences.append(translated_sentence)
                pbar.update(1)
            else:
                logger.error(f"Translation failed for sentence: {sentence}")
                return None
    return translated_sentences


def translate_text(
    text: str, target_language: str = "de", doc: Optional[Document] = None
) -> Union[Document, str]:
    """Main function to split, translate, and combine text."""
    logger.info(
        f"Starting translation process for text of length {len(text)} to {target_language}"
    )
    if doc:
        return translate_word_document(doc, target_language)

    sentences = split_into_sentences(text)
    translated_sentences = translate_sentences(sentences, target_language)
    if translated_sentences:
        final_translation = " ".join(translated_sentences)
        logger.info(f"Translation completed. Final length: {len(final_translation)}")
        return final_translation
    logger.error("Translation process failed")
    return None


def translate_word_document(doc, target_language="de"):
    """Translate a Word document while preserving its formatting."""
    logger.info("Starting Word document translation")

    # Create a new document for the translation
    translated_doc = Document()

    # Copy section properties and translate headers/footers
    for section in doc.sections:
        new_section = translated_doc.sections[-1]
        new_section._sectPr = section._sectPr

        # Translate header
        if section.header.is_linked_to_previous:
            new_section.header.is_linked_to_previous = True
        else:
            for paragraph in section.header.paragraphs:
                new_paragraph = new_section.header.paragraphs[-1]
                translate_paragraph(paragraph, new_paragraph, target_language)

        # Translate footer
        if section.footer.is_linked_to_previous:
            new_section.footer.is_linked_to_previous = True
        else:
            for paragraph in section.footer.paragraphs:
                new_paragraph = new_section.footer.paragraphs[-1]
                translate_paragraph(paragraph, new_paragraph, target_language)

    # Translate main document paragraphs
    for paragraph in doc.paragraphs:
        new_paragraph = translated_doc.add_paragraph()
        translate_paragraph(paragraph, new_paragraph, target_language)

    return translated_doc


def translate_paragraph(source_paragraph, target_paragraph, target_language):
    """Helper function to translate a paragraph while preserving formatting."""
    # Copy paragraph style
    target_paragraph.style = source_paragraph.style

    # Skip empty paragraphs
    if not source_paragraph.text.strip():
        return

    # Translate and preserve formatting for each run
    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(
            translate_text(run.text, target_language) if run.text.strip() else run.text
        )

        # Copy run formatting
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font.name:
            new_run.font.name = run.font.name
        if run.font.size:
            new_run.font.size = run.font.size
        if run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
